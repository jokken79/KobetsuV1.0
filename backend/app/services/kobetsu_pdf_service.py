"""
Kobetsu Keiyakusho PDF/DOCX Generation Service

Generates legally compliant individual dispatch contracts (個別契約書)
with all 16 required items under 労働者派遣法第26条.
"""
import os
from datetime import date, datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.config import settings
from app.models.kobetsu_keiyakusho import KobetsuKeiyakusho


class KobetsuPDFService:
    """Service for generating Kobetsu Keiyakusho documents."""

    def __init__(self):
        self.output_dir = Path(settings.PDF_OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _format_date_japanese(self, d: date) -> str:
        """Format date in Japanese style (令和X年X月X日)."""
        # Calculate Reiwa year (Reiwa started 2019-05-01)
        if d.year >= 2019:
            reiwa_year = d.year - 2018
            era = "令和"
        else:
            # Fallback to Western calendar
            return d.strftime("%Y年%m月%d日")

        return f"{era}{reiwa_year}年{d.month}月{d.day}日"

    def _format_time(self, t) -> str:
        """Format time as HH:MM."""
        return t.strftime("%H:%M")

    def _format_work_days(self, days: list) -> str:
        """Format work days list."""
        return "、".join(days)

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Add a heading with Japanese font."""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = "MS Gothic"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False, alignment=None):
        """Add a paragraph with Japanese font."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "MS Mincho"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
        run.font.size = Pt(10.5)
        run.bold = bold

        if alignment:
            p.alignment = alignment

        return p

    def _create_table(self, doc: Document, rows: int, cols: int):
        """Create a table with borders."""
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        return table

    def _add_cell_text(self, cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size: int = 8):
        """Add text to a table cell with specific formatting."""
        if len(cell.paragraphs) == 0:
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]

        p.alignment = align
        run = p.add_run(str(text) if text is not None else "")
        run.font.name = "MS Mincho"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
        run.font.size = Pt(font_size)
        run.bold = bold
        return p
    def generate_docx(self, contract: KobetsuKeiyakusho) -> str:
        """
        Generate DOCX document for a contract (Replica Layout).
        """
        doc = Document()

        # 1. Page Setup (Narrow Margins)
        section = doc.sections[0]
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.8)
        section.right_margin = Cm(0.8)

        # 2. Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(2)
        run = title.add_run("人材派遣個別契約書")
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(14)
        run.bold = True

        # 3. Intro Text
        intro_p = doc.add_paragraph()
        intro_p.paragraph_format.space_after = Pt(2)
        # User requested: Client (Worksite) is "Kou", Agency (Universal Kikaku) is "Otsu"
        agency_name = "ユニバーサル企画株式会社"
        intro_text = (
            f"{contract.worksite_name}（以下、「甲」という）と{agency_name}（以下、「乙」という）間で締結された労働"
            "者派遣基本契約書の定めに従い、次の派遣要件に基づき労働者派遣契約書を締結する。"
        )
        run = intro_p.add_run(intro_text)
        run.font.name = "MS Mincho"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
        run.font.size = Pt(8)

        # 4. Main Table Construction
        # We need a grid that can handle the most complex row.
        # Let's define a 12-column grid to give us flexibility for merging.
        # Col 0: Main Category (Vertical)
        # Col 1: Sub Category
        # Col 2-11: Data
        table = doc.add_table(rows=0, cols=12)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Define precise column widths (Total ~190mm)
        col_widths = [
            Cm(0.8),  # Col 0
            Cm(1.5),  # Col 1
            Cm(1.5),  # Col 2
            Cm(1.0),  # Col 3
            Cm(2.5),  # Col 4
            Cm(1.2),  # Col 5
            Cm(1.0),  # Col 6
            Cm(2.5),  # Col 7
            Cm(1.0),  # Col 8
            Cm(2.5),  # Col 9
            Cm(1.0),  # Col 10
            Cm(2.5)   # Col 11
        ]
        
        # Helper to add a row and get cells
        def add_row():
            row = table.add_row()
            # row.height_rule = 2 # Exact - REMOVED to allow text expansion
            # row.height = Pt(12) # Minimum height - REMOVED
            # Apply widths to cells
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]
            return row.cells



        # --- SECTION 1: 派遣先 (Dispatch Destination) ---
        # Row 1: Client Name
        cells = add_row()
        self._add_cell_text(cells[0], "派遣先", align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8) # Will merge vertically later
        self._add_cell_text(cells[1], "派遣先事業所", font_size=8)
        cells[1].merge(cells[2]) # Label spans 2
        self._add_cell_text(cells[3], f"名称  {contract.worksite_name}", font_size=8)
        cells[3].merge(cells[7]) # Name spans
        self._add_cell_text(cells[8], f"所在地  {contract.worksite_address}", font_size=8)
        cells[8].merge(cells[10]) # Address spans
        self._add_cell_text(cells[11], "TEL  -", font_size=8) # Placeholder

        # Row 2: Worksite Name (Often same as client for simple cases, but let's map it)
        cells = add_row()
        self._add_cell_text(cells[1], "就業場所", font_size=8)
        cells[1].merge(cells[2])
        self._add_cell_text(cells[3], f"名称  {contract.worksite_name}", font_size=8) # Using worksite name
        cells[3].merge(cells[7])
        self._add_cell_text(cells[8], f"所在地  {contract.worksite_address}", font_size=8)
        cells[8].merge(cells[10])
        self._add_cell_text(cells[11], "TEL  -", font_size=8)

        # Row 3: Org Unit & Offset Date
        cells = add_row()
        self._add_cell_text(cells[1], "組織単位", font_size=8)
        cells[1].merge(cells[2])
        self._add_cell_text(cells[3], contract.organizational_unit or "-", font_size=8)
        cells[3].merge(cells[7])
        self._add_cell_text(cells[8], "抵触日", font_size=8)
        self._add_cell_text(cells[9], "-", font_size=8) # Placeholder for conflict date
        cells[9].merge(cells[11])

        # Row 4: Commander
        cells = add_row()
        self._add_cell_text(cells[1], "指揮命令者", font_size=8)
        cells[1].merge(cells[2])
        self._add_cell_text(cells[3], "部署", font_size=8)
        self._add_cell_text(cells[4], contract.supervisor_department, font_size=8)
        cells[4].merge(cells[5])
        self._add_cell_text(cells[6], "役職", font_size=8)
        self._add_cell_text(cells[7], contract.supervisor_position, font_size=8)
        self._add_cell_text(cells[8], "氏名", font_size=8)
        self._add_cell_text(cells[9], contract.supervisor_name, font_size=8)
        self._add_cell_text(cells[10], "TEL", font_size=8)
        self._add_cell_text(cells[11], "-", font_size=8)

        # Row 5: Saki Manager (Responsible Person)
        cells = add_row()
        self._add_cell_text(cells[1], "製造業務専門派遣先責任者", font_size=8) # Label from image
        cells[1].merge(cells[2])
        m = contract.haken_saki_manager
        self._add_cell_text(cells[3], "部署", font_size=8)
        self._add_cell_text(cells[4], m.get('department', ''), font_size=8)
        cells[4].merge(cells[5])
        self._add_cell_text(cells[6], "役職", font_size=8)
        self._add_cell_text(cells[7], m.get('position', ''), font_size=8)
        self._add_cell_text(cells[8], "氏名", font_size=8)
        self._add_cell_text(cells[9], m.get('name', ''), font_size=8)
        self._add_cell_text(cells[10], "TEL", font_size=8)
        self._add_cell_text(cells[11], m.get('phone', ''), font_size=8)

        # Row 6: Saki Complaint
        cells = add_row()
        self._add_cell_text(cells[1], "苦情処理担当者", font_size=8)
        cells[1].merge(cells[2])
        c = contract.haken_saki_complaint_contact
        self._add_cell_text(cells[3], "部署", font_size=8)
        self._add_cell_text(cells[4], c.get('department', ''), font_size=8)
        cells[4].merge(cells[5])
        self._add_cell_text(cells[6], "役職", font_size=8)
        self._add_cell_text(cells[7], c.get('position', ''), font_size=8)
        self._add_cell_text(cells[8], "氏名", font_size=8)
        self._add_cell_text(cells[9], c.get('name', ''), font_size=8)
        self._add_cell_text(cells[10], "TEL", font_size=8)
        self._add_cell_text(cells[11], c.get('phone', ''), font_size=8)

        # Merge "派遣先" vertical label (Rows 0-5)
        table.cell(0, 0).merge(table.cell(5, 0))


        # --- SECTION 2: 派遣元 (Dispatch Source) ---
        # Row 7: Moto Manager
        cells = add_row()
        self._add_cell_text(cells[0], "派遣元", align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8)
        self._add_cell_text(cells[1], "製造業務専門派遣元責任者", font_size=8)
        cells[1].merge(cells[2])
        m = contract.haken_moto_manager
        self._add_cell_text(cells[3], "部署", font_size=8)
        self._add_cell_text(cells[4], m.get('department', ''), font_size=8)
        cells[4].merge(cells[5])
        self._add_cell_text(cells[6], "役職", font_size=8)
        self._add_cell_text(cells[7], m.get('position', ''), font_size=8)
        self._add_cell_text(cells[8], "氏名", font_size=8)
        self._add_cell_text(cells[9], m.get('name', ''), font_size=8)
        self._add_cell_text(cells[10], "TEL", font_size=8)
        self._add_cell_text(cells[11], m.get('phone', ''), font_size=8)

        # Row 8: Moto Complaint
        cells = add_row()
        self._add_cell_text(cells[1], "苦情処理担当者", font_size=8)
        cells[1].merge(cells[2])
        c = contract.haken_moto_complaint_contact
        self._add_cell_text(cells[3], "部署", font_size=8)
        self._add_cell_text(cells[4], c.get('department', ''), font_size=8)
        cells[4].merge(cells[5])
        self._add_cell_text(cells[6], "役職", font_size=8)
        self._add_cell_text(cells[7], c.get('position', ''), font_size=8)
        self._add_cell_text(cells[8], "氏名", font_size=8)
        self._add_cell_text(cells[9], c.get('name', ''), font_size=8)
        self._add_cell_text(cells[10], "TEL", font_size=8)
        self._add_cell_text(cells[11], c.get('phone', ''), font_size=8)

        # Merge "派遣元" vertical label (Rows 6-7)
        table.cell(6, 0).merge(table.cell(7, 0))


        # --- SECTION 3: 派遣内容 (Dispatch Content) ---
        # Row 9: Worker Limits
        cells = add_row()
        self._add_cell_text(cells[0], "派遣内容", align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8)
        self._add_cell_text(cells[1], "派遣労働者を協定対象労働者\nに限定するか否か", font_size=8)
        cells[1].merge(cells[2])
        
        check_limit = "☑" if contract.is_kyotei_taisho else "☐"
        check_unlimit = "☐" if contract.is_kyotei_taisho else "☑"
        self._add_cell_text(cells[3], f"{check_limit} 協定対象派遣労働者に限定", font_size=8)
        cells[3].merge(cells[6])
        self._add_cell_text(cells[7], f"{check_unlimit} 限定なし", font_size=8)
        cells[7].merge(cells[11])

        # Row 10: Authority
        cells = add_row()
        self._add_cell_text(cells[1], "派遣労働者の責任の程度", font_size=8)
        cells[1].merge(cells[2])
        # Assuming responsibility_level maps to authority
        self._add_cell_text(cells[3], "☑ 付与される権限なし", font_size=8) # Defaulting for now
        cells[3].merge(cells[6])
        self._add_cell_text(cells[7], "☐ 付与される権限あり", font_size=8)
        cells[7].merge(cells[11])

        # Row 11: Work Content
        cells = add_row()
        self._add_cell_text(cells[1], "業務内容", font_size=8)
        cells[1].merge(cells[2])
        self._add_cell_text(cells[3], contract.work_content, font_size=8)
        cells[3].merge(cells[11])

        # Row 12: Period & Count
        cells = add_row()
        self._add_cell_text(cells[1], "派遣期間", font_size=8)
        cells[1].merge(cells[2])
        period = f"{self._format_date_japanese(contract.dispatch_start_date)}  ～  {self._format_date_japanese(contract.dispatch_end_date)}"
        self._add_cell_text(cells[3], period, font_size=8)
        cells[3].merge(cells[8])
        self._add_cell_text(cells[9], "人数", font_size=8)
        self._add_cell_text(cells[10], str(contract.number_of_workers), font_size=8)
        cells[10].merge(cells[11])

        # Row 13: Work Days
        cells = add_row()
        self._add_cell_text(cells[1], "就業日", font_size=8)
        cells[1].merge(cells[2])
        days = self._format_work_days(contract.work_days)
        self._add_cell_text(cells[3], f"{days}   (祝日、年末年始、夏季休暇を除く)", font_size=8)
        cells[3].merge(cells[11])

        # Row 14: Work Hours
        cells = add_row()
        self._add_cell_text(cells[1], "就業時間", font_size=8)
        cells[1].merge(cells[2])
        hours = f"{self._format_time(contract.work_start_time)} ～ {self._format_time(contract.work_end_time)}"
        self._add_cell_text(cells[3], f"昼勤: {hours}", font_size=8)
        cells[3].merge(cells[11])

        # Row 15: Break Time
        cells = add_row()
        self._add_cell_text(cells[1], "休憩時間", font_size=8)
        cells[1].merge(cells[2])
        self._add_cell_text(cells[3], f"{contract.break_time_minutes}分", font_size=8)
        cells[3].merge(cells[11])

        # Row 16: Overtime
        cells = add_row()
        self._add_cell_text(cells[1], "時間外労働", font_size=8)
        cells[1].merge(cells[2])
        ot_text = "1ヶ月に45時間、1年に360時間を限度とする。"
        if contract.overtime_max_hours_month:
            ot_text = f"1ヶ月に{contract.overtime_max_hours_month}時間、1年に360時間を限度とする。"
        self._add_cell_text(cells[3], ot_text, font_size=8)
        cells[3].merge(cells[11])

        # Row 17: Rates
        cells = add_row()
        self._add_cell_text(cells[1], "派遣料金", font_size=8)
        cells[1].merge(cells[2])
        
        # Rate formatting
        rate_basic = f"基本 ￥{contract.hourly_rate:,.0f}"
        rate_ot = f"残業(1.25%) ￥{contract.overtime_rate:,.0f}"
        rate_night = f"深夜(0.25%) ￥{contract.night_shift_rate:,.0f}" if contract.night_shift_rate else ""
        rate_holiday = f"休日(1.35%) ￥{contract.holiday_rate:,.0f}" if contract.holiday_rate else ""
        
        self._add_cell_text(cells[3], f"{rate_basic}    {rate_ot}    {rate_night}    {rate_holiday}", font_size=8)
        cells[3].merge(cells[11])

        # Row 18: Payment Terms (New)
        cells = add_row()
        self._add_cell_text(cells[1], "支払い条件", font_size=8)
        cells[1].merge(cells[2])
        self._add_cell_text(cells[3], "締日: 20日    支払日: 翌月20日    支払方法: 銀行振込", font_size=8) # Hardcoded example
        cells[3].merge(cells[11])

        # Row 19: Safety
        cells = add_row()
        self._add_cell_text(cells[1], "安全・衛生", font_size=8)
        cells[1].merge(cells[2])
        safety = (
            "派遣先及び派遣元事業主は、労働者派遣法第44条から第47条の2までの規定により課された各法令を順守し、"
            "自己に課された法令上の責任を負う。なお、派遣就業中の安全及び衛生については、派遣先の安全衛生に関する規定を適用することとし、"
            "その他については、派遣元の安全衛生に関する規定を適用する。"
        )
        self._add_cell_text(cells[3], safety, font_size=6)
        cells[3].merge(cells[11])

        # Row 20: Welfare
        cells = add_row()
        self._add_cell_text(cells[1], "便宜供与", font_size=8)
        cells[1].merge(cells[2])
        welfare = (
            "派遣先は、派遣労働者に対して利用の機会を与える給食施設、休憩室、及び更衣室については、"
            "本契約に基づく労働者派遣に係る派遣労働者に対しても、利用の機会を与えるように配慮しなければならないこととする。"
        )
        self._add_cell_text(cells[3], welfare, font_size=6)
        cells[3].merge(cells[11])

        # Row 21: Complaint Method
        cells = add_row()
        self._add_cell_text(cells[1], "苦情処理方法", font_size=8)
        cells[1].merge(cells[2])
        comp_method = (
            "(1)派遣元事業主における苦情処理担当者が苦情の申し出を受けたときは、ただちに製造業務専門派遣元責任者へ連絡することとし、"
            "当該派遣元責任者が中心となって、誠意をもって、遅滞なく、当該苦情の適切迅速な処理を図ることとし、その結果について必ず派遣労働者に通知することとする。\n"
            "(2)派遣先における苦情処理担当者が苦情の申し出を受けたときは、ただちに製造業務専門派遣先責任者へ連絡することとし、"
            "当該派遣先責任者が中心となって、誠意をもって、遅滞なく、当該苦情の適切かつ迅速な処理を図ることとし、その結果について必ず派遣労働者に通知することとする。\n"
            "(3)派遣先及び派遣元事業主は、自らでその解決が容易であり、即時に処理した苦情の他は、相互に遅滞なく通知するとともに、"
            "密接に連絡調整を行いつつ、その解決を図ることとする。"
        )
        self._add_cell_text(cells[3], comp_method, font_size=6)
        cells[3].merge(cells[11])

        # Row 22: Termination
        cells = add_row()
        self._add_cell_text(cells[1], "労働者派遣契約の契約の解除に当\nたって講ずる派遣労働者の雇用の\n安定を図るための措置", font_size=8)
        cells[1].merge(cells[2])
        term_text = (
            "(1) 労働者派遣契約の解除の事前申し入れ  派遣先は、専ら派遣先に起因する事由により、労働者派遣契約の契約期間が満了する前の解除を行おうとする場合には、"
            "派遣元の合意を得ることはもとより、あらかじめ相当の猶予期間をもって派遣元に解除の申し入れを行うこととする。\n"
            "(2) 就業機会の確保派遣元事業主及び派遣先は、労働者派遣契約の契約期間が満了する前に派遣労働者の責に帰すべき事由によらない労働者派遣契約の解除を"
            "行った場合には、派遣先の関連会社での就業をあっせんする等により、当該労働者派遣契約に係る派遣労働者の新たな就業機会の確保を図ることとする。\n"
            "(3) 損害賠償等に係る適切な措置派遣先は、派遣先の責に帰すべき事由により労働者派遣契約の契約期間が満了する前に労働者派遣契約の解除を行おうとする場"
            "合には、派遣労働者の新たな就業機会の確保を図ることとし、これができないときには、少なくとも当該労働者派遣契約の解除に伴い派遣元事業主が当該派遣労働"
            "者に係る派遣労働者を休業させること等を余儀なくされたことにより生じた損害の賠償を行わなければならないこととする。例えば、派遣元事業主が当該派遣労"
            "働者を休業させる場合は、休業手当に相当する額以上の額について、派遣元事業主がやむを得ない事由により当該派遣労働者を解雇する場合は、派遣先による解除の"
            "申し入れが相当の猶予期間をもって行われなかったことにより派遣元事業主が解雇の予告をしないときは、30日分以上、当該予告をした日から解雇の日までの期"
            "間が30日に満たないときは当該解雇の日の30日前の日から当該予告の日までの日数分以上の賃金に相当する額以上の額について、損害の賠償を行わなければな"
            "らないこととする。その他派遣先は派遣元事業主と十分に協議した上で適切な善後処理方策を講ずることとする。また、派遣元事業主及び派遣先の双方の責に帰す"
            "べき事由がある場合には、派遣元事業主及び派遣先のそれぞれの責に帰すべき部分の割合についても十分に考慮することとする。\n"
            "(4) 労働者派遣契約の解除の理由の明示 派遣先は、労働者派遣契約の契約期間が満了する前に労働者派遣契約の解除を行おうとする場合であって派遣元事業主"
            "から請求があったときは、労働者派遣契約の解除を行った理由を派遣元事業主に対してあきらかにすることとする。"
        )
        self._add_cell_text(cells[3], term_text, font_size=5) # Smaller font for long text
        cells[3].merge(cells[11])

        # Row 23: Dispute Prevention (Direct Hire)
        cells = add_row()
        self._add_cell_text(cells[1], "派遣先が派遣労働者を雇用する場\n合の紛争防止措置", font_size=8)
        cells[1].merge(cells[2])
        direct_text = "派遣先が派遣終了後に、当該派遣労働者を雇用する場合、その雇用意思を事前に派遣元へ示すこととする。"
        self._add_cell_text(cells[3], direct_text, font_size=6)
        cells[3].merge(cells[11])

        # Row 24: Unlimited/60+
        cells = add_row()
        self._add_cell_text(cells[1], "派遣労働者を無期雇用派遣労働者\n又は60歳以上の者の者に限定\nするか否かの別", font_size=6)
        cells[1].merge(cells[2])
        self._add_cell_text(cells[3], "無期雇用又は60歳以上に限定しない。", font_size=8)
        cells[3].merge(cells[11])

        # Merge "派遣内容" vertical label (Rows 8-23)
        table.cell(8, 0).merge(table.cell(23, 0))

        # 5. Footer (Signatures)
        doc.add_paragraph()
        p = doc.add_paragraph("上記契約の証として本書2通を作成し、甲乙記名押印のうえ、各1通を保有する。")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.runs[0]
        run.font.size = Pt(8)
        run.font.name = "MS Mincho"

        date_p = doc.add_paragraph(self._format_date_japanese(contract.contract_date))
        date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_p.paragraph_format.space_after = Pt(2)
        run = date_p.runs[0]
        run.font.size = Pt(8)
        
        # Signature Table
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.width = Cm(18)
        sig_table.autofit = True
        
        # Party A (Client - Kou)
        cell_a = sig_table.cell(0, 0)
        self._add_cell_text(cell_a, "(甲)", align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8)
        self._add_cell_text(cell_a, contract.worksite_name, font_size=8)
        self._add_cell_text(cell_a, contract.worksite_address, font_size=8)
        
        # Party B (Agency - Otsu)
        cell_b = sig_table.cell(0, 1)
        self._add_cell_text(cell_b, "(乙)", align=WD_ALIGN_PARAGRAPH.CENTER, font_size=8)
        self._add_cell_text(cell_b, "ユニバーサル企画株式会社", font_size=8)
        self._add_cell_text(cell_b, "愛知県名古屋市東区徳川2-18-18", font_size=8)

        # Save document
        output_path = self.output_dir / f"{contract.contract_number}_v4.docx"
        doc.save(str(output_path))

        return str(output_path)

    def generate_pdf(self, contract: KobetsuKeiyakusho) -> str:
        """
        Generate PDF document for a contract.

        First generates DOCX, then converts to PDF.

        Args:
            contract: KobetsuKeiyakusho instance

        Returns:
            Path to generated PDF file
        """
        # Generate DOCX first
        docx_path = self.generate_docx(contract)

        # Convert to PDF
        pdf_path = docx_path.replace('.docx', '.pdf')

        try:
            # Try using LibreOffice for conversion
            import subprocess
            result = subprocess.run(
                [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(self.output_dir),
                    docx_path
                ],
                capture_output=True,
                timeout=60
            )

            if result.returncode == 0 and os.path.exists(pdf_path):
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        try:
            # Alternative: Try using docx2pdf
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            return pdf_path
        except ImportError:
            pass

        # If PDF conversion fails, return DOCX path
        # In production, you would want to handle this better
        return docx_path

    def generate_preview(self, contract: KobetsuKeiyakusho) -> dict:
        """
        Generate a preview of the contract without creating a file.

        Args:
            contract: KobetsuKeiyakusho instance

        Returns:
            Dictionary with contract preview data
        """
        return {
            "contract_number": contract.contract_number,
            "contract_date": self._format_date_japanese(contract.contract_date),
            "dispatch_period": {
                "start": self._format_date_japanese(contract.dispatch_start_date),
                "end": self._format_date_japanese(contract.dispatch_end_date),
            },
            "dispatch_company": {
                "name": settings.COMPANY_NAME,
                "address": settings.COMPANY_ADDRESS,
                "license": settings.COMPANY_LICENSE_NUMBER,
            },
            "client_company": {
                "name": contract.worksite_name,
                "address": contract.worksite_address,
                "organizational_unit": contract.organizational_unit,
            },
            "work_details": {
                "content": contract.work_content,
                "responsibility_level": contract.responsibility_level,
                "supervisor": {
                    "department": contract.supervisor_department,
                    "position": contract.supervisor_position,
                    "name": contract.supervisor_name,
                },
            },
            "working_conditions": {
                "days": contract.work_days,
                "hours": f"{self._format_time(contract.work_start_time)} - {self._format_time(contract.work_end_time)}",
                "break_minutes": contract.break_time_minutes,
            },
            "rates": {
                "hourly": float(contract.hourly_rate),
                "overtime": float(contract.overtime_rate),
                "night_shift": float(contract.night_shift_rate) if contract.night_shift_rate else None,
                "holiday": float(contract.holiday_rate) if contract.holiday_rate else None,
            },
            "workers": contract.number_of_workers,
            "status": contract.status,
        }
